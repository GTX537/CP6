from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


ROOT = Path(__file__).resolve().parent
SOURCE_MD = ROOT / "販売管理_ユーザー操作マニュアル.md"
OUTPUT_DOCX = ROOT / "销售管理用户操作手册.docx"

ACCENT = RGBColor(31, 78, 121)
ACCENT_LIGHT = "DCE6F1"
HEADER_FILL = "B4C6E7"
GRID = "A6A6A6"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def ensure_rfonts(run) -> OxmlElement:
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.rFonts
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.append(r_fonts)
    return r_fonts


def ensure_style_rfonts(style) -> OxmlElement:
    r_pr = style._element.get_or_add_rPr()
    r_fonts = r_pr.rFonts
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.append(r_fonts)
    return r_fonts


def set_table_borders(table) -> None:
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        elem = OxmlElement(f"w:{edge}")
        elem.set(qn("w:val"), "single")
        elem.set(qn("w:sz"), "6")
        elem.set(qn("w:color"), GRID)
        borders.append(elem)
    tbl_pr.append(borders)


def set_page_margins(section) -> None:
    section.top_margin = Cm(1.8)
    section.bottom_margin = Cm(1.8)
    section.left_margin = Cm(2.0)
    section.right_margin = Cm(2.0)
    section.header_distance = Cm(0.8)
    section.footer_distance = Cm(0.8)


def set_default_fonts(document: Document) -> None:
    styles = document.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10.5)
    ensure_style_rfonts(normal).set(qn("w:eastAsia"), "Microsoft YaHei")

    for style_name in ("Title", "Subtitle", "Heading 1", "Heading 2", "Heading 3"):
        style = styles[style_name]
        style.font.name = "Calibri"
        ensure_style_rfonts(style).set(qn("w:eastAsia"), "Microsoft YaHei")

    styles["Title"].font.size = Pt(24)
    styles["Title"].font.bold = True
    styles["Title"].font.color.rgb = ACCENT

    styles["Subtitle"].font.size = Pt(11)
    styles["Subtitle"].font.color.rgb = RGBColor(90, 90, 90)

    styles["Heading 1"].font.size = Pt(16)
    styles["Heading 1"].font.bold = True
    styles["Heading 1"].font.color.rgb = ACCENT

    styles["Heading 2"].font.size = Pt(12.5)
    styles["Heading 2"].font.bold = True
    styles["Heading 2"].font.color.rgb = RGBColor(47, 84, 150)

    styles["Heading 3"].font.size = Pt(11.5)
    styles["Heading 3"].font.bold = True
    styles["Heading 3"].font.color.rgb = RGBColor(68, 114, 196)


def add_header_footer(section) -> None:
    header = section.header
    hp = header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = hp.add_run("CP6 销售管理用户操作手册")
    run.font.size = Pt(8.5)
    run.font.color.rgb = RGBColor(120, 120, 120)
    ensure_rfonts(run).set(qn("w:eastAsia"), "Microsoft YaHei")

    footer = section.footer
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = fp.add_run("第 ")
    run.font.size = Pt(8.5)
    ensure_rfonts(run).set(qn("w:eastAsia"), "Microsoft YaHei")
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), "PAGE")
    r = OxmlElement("w:r")
    t = OxmlElement("w:t")
    t.text = "1"
    r.append(t)
    fld.append(r)
    fp._p.append(fld)
    run = fp.add_run(" 页")
    run.font.size = Pt(8.5)
    ensure_rfonts(run).set(qn("w:eastAsia"), "Microsoft YaHei")


def add_cover(document: Document) -> None:
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.space_before = Pt(60)
    p.space_after = Pt(8)
    r = p.add_run("CP6")
    r.font.size = Pt(18)
    r.font.bold = True
    r.font.color.rgb = ACCENT
    ensure_rfonts(r).set(qn("w:eastAsia"), "Microsoft YaHei")

    p = document.add_paragraph(style="Title")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.space_after = Pt(6)
    p.add_run("销售管理用户操作手册")

    p = document.add_paragraph(style="Subtitle")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run("适用于销售、营业助理、报价担当、受注担当及主数据维护人员")

    table = document.add_table(rows=4, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    set_table_borders(table)
    labels = ["系统名称", "文档用途", "更新日期", "说明"]
    values = [
        "CP6 销售管理",
        "用户培训、上线说明、日常操作参考",
        "2026-05-08",
        "本手册依据当前已开发完成的系统功能整理。",
    ]
    widths = [Cm(4.0), Cm(10.5)]
    for row_idx, (label, value) in enumerate(zip(labels, values)):
        for col_idx, text in enumerate((label, value)):
            cell = table.cell(row_idx, col_idx)
            cell.width = widths[col_idx]
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.left_indent = Pt(2)
            run = p.add_run(text)
            run.font.size = Pt(10.5)
            ensure_rfonts(run).set(qn("w:eastAsia"), "Microsoft YaHei")
            if col_idx == 0:
                run.bold = True
                set_cell_shading(cell, ACCENT_LIGHT)

    document.add_paragraph().add_run().add_break(WD_BREAK.PAGE)


def add_intro_box(document: Document) -> None:
    p = document.add_paragraph()
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run("使用说明")
    run.bold = True
    run.font.size = Pt(12)
    run.font.color.rgb = ACCENT
    ensure_rfonts(run).set(qn("w:eastAsia"), "Microsoft YaHei")

    table = document.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    set_table_borders(table)
    cell = table.cell(0, 0)
    cell.width = Cm(16.2)
    set_cell_shading(cell, "F7FBFF")
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.space_before = Pt(0)
    for text in [
        "本手册按实际业务顺序编排，建议培训时优先演示“交易先 -> 报价计算 -> 正式报价 -> 产品 -> 受注”。",
        "如遇检索无结果、必填报错或并发提示，请先检查检索条件、必填项及是否需要重新读取数据。",
    ]:
        r = p.add_run(f"• {text}\n")
        r.font.size = Pt(10.5)
        ensure_rfonts(r).set(qn("w:eastAsia"), "Microsoft YaHei")


def apply_run_font(run, *, bold=False, color=None) -> None:
    run.font.name = "Calibri"
    ensure_rfonts(run).set(qn("w:eastAsia"), "Microsoft YaHei")
    run.font.size = Pt(10.5)
    run.bold = bold
    if color:
        run.font.color.rgb = color


def add_paragraph_with_inline_formatting(document: Document, text: str, style: str | None = None) -> None:
    p = document.add_paragraph(style=style)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.2

    parts = re.split(r"(`[^`]+`)", text)
    for part in parts:
        if not part:
            continue
        if part.startswith("`") and part.endswith("`"):
            run = p.add_run(part[1:-1])
            apply_run_font(run, bold=True, color=ACCENT)
        else:
            run = p.add_run(part)
            apply_run_font(run)


def add_bullet(document: Document, text: str) -> None:
    p = document.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = 1.15
    prefix = p.add_run("• ")
    apply_run_font(prefix, bold=True, color=ACCENT)
    parts = re.split(r"(`[^`]+`)", text)
    for part in parts:
        if not part:
            continue
        if part.startswith("`") and part.endswith("`"):
            run = p.add_run(part[1:-1])
            apply_run_font(run, bold=True, color=ACCENT)
        else:
            run = p.add_run(part)
            apply_run_font(run)

def add_numbered(document: Document, marker: str, text: str) -> None:
    p = document.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = 1.15
    prefix = p.add_run(f"{marker} ")
    apply_run_font(prefix, bold=True, color=ACCENT)
    parts = re.split(r"(`[^`]+`)", text)
    for part in parts:
        if not part:
            continue
        if part.startswith("`") and part.endswith("`"):
            run = p.add_run(part[1:-1])
            apply_run_font(run, bold=True, color=ACCENT)
        else:
            run = p.add_run(part)
            apply_run_font(run)


def add_table(document: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = document.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    set_table_borders(table)

    widths = []
    for idx, header in enumerate(headers):
        if idx == 0 and len(headers) >= 3:
            widths.append(Cm(2.2))
        elif len(header) <= 6:
            widths.append(Cm(3.2))
        else:
            widths.append(Cm(max(4.2, 16.0 / len(headers))))

    hdr_cells = table.rows[0].cells
    for idx, head in enumerate(headers):
        cell = hdr_cells[idx]
        cell.width = widths[idx]
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        set_cell_shading(cell, HEADER_FILL)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run(head)
        apply_run_font(run, bold=True)

    for row in rows:
        row_cells = table.add_row().cells
        for idx, value in enumerate(row):
            cell = row_cells[idx]
            cell.width = widths[idx]
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.1
            if len(value) <= 12:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            else:
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            run = p.add_run(value)
            apply_run_font(run)

    document.add_paragraph()


def parse_markdown_tables(lines: list[str], start: int):
    headers = [part.strip() for part in lines[start].strip().strip("|").split("|")]
    rows = []
    idx = start + 2
    while idx < len(lines):
        line = lines[idx].rstrip()
        if not line.startswith("|"):
            break
        rows.append([part.strip() for part in line.strip().strip("|").split("|")])
        idx += 1
    return headers, rows, idx


def build_docx() -> Path:
    text = SOURCE_MD.read_text(encoding="utf-8")
    lines = text.splitlines()

    document = Document()
    for section in document.sections:
        set_page_margins(section)
        add_header_footer(section)
    set_default_fonts(document)
    add_cover(document)
    add_intro_box(document)

    idx = 0
    while idx < len(lines):
        raw = lines[idx].rstrip()
        line = raw.strip()
        if not line or line == "---":
            idx += 1
            continue
        if line.startswith(">"):
            idx += 1
            continue
        if line.startswith("# "):
            idx += 1
            continue
        if line.startswith("## "):
            add_paragraph_with_inline_formatting(document, line[3:].strip(), style="Heading 1")
            idx += 1
            continue
        if line.startswith("### "):
            add_paragraph_with_inline_formatting(document, line[4:].strip(), style="Heading 2")
            idx += 1
            continue
        if re.match(r"^\|.+\|$", line) and idx + 1 < len(lines) and re.match(r"^\|(?:\s*:?-+:?\s*\|)+$", lines[idx + 1].strip()):
            headers, rows, idx = parse_markdown_tables(lines, idx)
            add_table(document, headers, rows)
            continue
        if re.match(r"^\d+\.\s+", line):
            marker = re.match(r"^(\d+\.)\s+", line).group(1)
            add_numbered(document, marker, re.sub(r"^\d+\.\s+", "", line))
            idx += 1
            continue
        if line.startswith("- "):
            add_bullet(document, line[2:].strip())
            idx += 1
            continue
        add_paragraph_with_inline_formatting(document, line)
        idx += 1

    section = document.sections[-1]
    section.start_type = WD_SECTION.NEW_PAGE
    document.save(OUTPUT_DOCX)
    return OUTPUT_DOCX


if __name__ == "__main__":
    path = build_docx()
    print(path)
